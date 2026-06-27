#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
TML Markdown Editor - 性能优化版
移除了 MathJax 和 Mermaid，添加预览防抖，大幅提升响应速度
保留悬停放大功能
"""

import sys
import os
import ctypes
import html as html_lib
import re
import importlib
import time
import traceback

# PyInstaller 隐藏导入标记：让静态分析能检测到懒加载的模块
# 运行时不会执行（if False），但 PyInstaller 会扫描到并打包这些依赖
if False:
    import markdown
    import markdown.extensions.extra
    import markdown.extensions.tables
    import markdown.extensions.fenced_code
    import docx
    from PyQt6.QtWebEngineWidgets import QWebEngineView
    from PyQt6.QtWebEngineCore import QWebEngineSettings
    from PyQt6.QtWebChannel import QWebChannel

from PyQt6.QtWidgets import (
    QApplication, QMainWindow, QSplitter, QTextEdit,
    QFileDialog, QMessageBox, QMenuBar, QMenu,
    QToolBar, QStatusBar, QWidget, QVBoxLayout,
    QTabWidget, QDialog, QPushButton, QSplashScreen
)
from PyQt6.QtGui import QAction, QFont, QSyntaxHighlighter, QTextCharFormat, QColor, QIcon, QTextCursor, QPixmap, QPainter, QLinearGradient
from PyQt6.QtCore import Qt, QFileInfo, QUrl, QObject, pyqtSignal, pyqtSlot, QTimer


_MARKDOWN_MODULE = None
_DOCX_DOCUMENT_CLASS = None
_DOCX_IMPORT_FAILED = False


# ==================== 启动日志模块 ====================
class StartupLogger:
    def __init__(self):
        self.started_at = time.perf_counter()
        self.log_path = self._resolve_log_path()
        self.enabled = False

    def _resolve_log_path(self):
        if getattr(sys, "frozen", False):
            base_dir = os.path.dirname(sys.executable)
        else:
            base_dir = os.path.dirname(os.path.abspath(__file__))
        return os.path.join(base_dir, "TMLEditor_startup.log")

    def _write(self, message):
        if not self.enabled:
            return
        elapsed_ms = (time.perf_counter() - self.started_at) * 1000.0
        line = f"[{elapsed_ms:9.1f} ms] {message}\n"
        try:
            with open(self.log_path, "a", encoding="utf-8") as log_file:
                log_file.write(line)
        except Exception:
            self.enabled = False

    def mark(self, message):
        self._write(message)

    def section(self, message):
        self._write(message)


STARTUP_LOGGER = StartupLogger()


def get_markdown_module():
    global _MARKDOWN_MODULE
    if _MARKDOWN_MODULE is None:
        _MARKDOWN_MODULE = importlib.import_module("markdown")
    return _MARKDOWN_MODULE


def get_docx_document_class():
    global _DOCX_DOCUMENT_CLASS, _DOCX_IMPORT_FAILED
    if _DOCX_DOCUMENT_CLASS is not None:
        return _DOCX_DOCUMENT_CLASS
    if _DOCX_IMPORT_FAILED:
        return None
    try:
        _DOCX_DOCUMENT_CLASS = importlib.import_module("docx").Document
        return _DOCX_DOCUMENT_CLASS
    except Exception:
        _DOCX_IMPORT_FAILED = True
        return None


def log_startup(message):
    STARTUP_LOGGER.mark(message)


def install_global_exception_hook():
    def hook(exc_type, exc_value, exc_traceback):
        try:
            log_startup("unhandled exception")
            log_startup("".join(traceback.format_exception(exc_type, exc_value, exc_traceback)).rstrip())
        finally:
            sys.__excepthook__(exc_type, exc_value, exc_traceback)

    sys.excepthook = hook


# ==================== 启动画面（Splash Screen） ====================
class SplashScreen(QSplashScreen):
    def __init__(self):
        pixmap = QPixmap(400, 280)
        pixmap.fill(QColor("#1e1e2e"))
        super().__init__(pixmap)
        self.setFixedSize(400, 280)
        self.setWindowFlag(Qt.WindowType.WindowStaysOnTopHint, True)

        self._progress = 0
        self._message = "正在启动..."
        self._dots = 0
        self._dot_timer = QTimer(self)
        self._dot_timer.timeout.connect(self._animate_dots)
        self._dot_timer.start(400)

        self._draw_content()

    def _animate_dots(self):
        self._dots = (self._dots + 1) % 4
        self._draw_content()

    def set_progress(self, value, message=None):
        self._progress = max(0, min(100, value))
        if message:
            self._message = message
        self._draw_content()

    def _draw_content(self):
        pixmap = QPixmap(400, 280)
        pixmap.fill(QColor("#1e1e2e"))
        painter = QPainter(pixmap)
        painter.setRenderHint(QPainter.RenderHint.Antialiasing, True)

        title_font = QFont("Microsoft YaHei", 22, QFont.Weight.Bold)
        painter.setFont(title_font)
        painter.setPen(QColor("#89b4fa"))
        painter.drawText(pixmap.rect(), Qt.AlignmentFlag.AlignHCenter | Qt.AlignmentFlag.AlignTop,
                         "\n\nTML Markdown")

        subtitle_font = QFont("Microsoft YaHei", 10)
        painter.setFont(subtitle_font)
        painter.setPen(QColor("#a6adc8"))
        painter.drawText(pixmap.rect(), Qt.AlignmentFlag.AlignHCenter | Qt.AlignmentFlag.AlignTop,
                         "\n\n\n\n\n-")

        dots = "." * self._dots
        msg_font = QFont("Microsoft YaHei", 9)
        painter.setFont(msg_font)
        painter.setPen(QColor("#cdd6f4"))
        painter.drawText(pixmap.rect(), Qt.AlignmentFlag.AlignHCenter | Qt.AlignmentFlag.AlignBottom,
                         f"\n\n{self._message}{dots}\n\n")

        bar_x = 50
        bar_y = 190
        bar_w = 300
        bar_h = 6

        painter.setPen(Qt.PenStyle.NoPen)
        painter.setBrush(QColor("#313244"))
        painter.drawRoundedRect(bar_x, bar_y, bar_w, bar_h, 3, 3)

        fill_w = int(bar_w * (self._progress / 100.0))
        if fill_w > 0:
            gradient = QLinearGradient(bar_x, 0, bar_x + bar_w, 0)
            gradient.setColorAt(0, QColor("#89b4fa"))
            gradient.setColorAt(1, QColor("#cba6f7"))
            painter.setBrush(gradient)
            painter.drawRoundedRect(bar_x, bar_y, fill_w, bar_h, 3, 3)

        painter.end()
        self.setPixmap(pixmap)


# ==================== Markdown 语法高亮器（优化版） ====================
class MarkdownHighlighter(QSyntaxHighlighter):
    # 类常量：限制高亮的最大行数，避免大文件卡顿
    MAX_HIGHLIGHT_LINES = 1000

    def __init__(self, parent=None):
        super().__init__(parent)
        self.rules = []
        # 预编译正则表达式以提升性能
        self._compile_patterns()

    def _compile_patterns(self):
        # 标题 (# ...)
        title_format = QTextCharFormat()
        title_format.setForeground(QColor(0, 120, 200))
        title_format.setFontWeight(QFont.Weight.Bold)
        self.rules.append((re.compile(r"^#{1,6}\s+.*"), title_format))

        # 粗体 (**bold**)
        bold_format = QTextCharFormat()
        bold_format.setFontWeight(QFont.Weight.Bold)
        self.rules.append((re.compile(r"\*\*[^*]+\*\*"), bold_format))

        # 斜体 (*italic*)
        italic_format = QTextCharFormat()
        italic_format.setFontItalic(True)
        self.rules.append((re.compile(r"(?<!\*)\*[^*]+\*(?!\*)"), italic_format))

        # 行内代码 (`code`)
        code_format = QTextCharFormat()
        code_format.setForeground(QColor(150, 100, 50))
        code_format.setFont(QFont("Consolas"))
        self.rules.append((re.compile(r"`[^`]+`"), code_format))

        # 链接 [text](url)
        link_format = QTextCharFormat()
        link_format.setForeground(QColor(0, 150, 0))
        link_format.setFontUnderline(True)
        self.rules.append((re.compile(r"\[[^\]]+\]\([^\)]+\)"), link_format))

    def highlightBlock(self, text):
        # 限制高亮处理的行数，避免大文件卡顿
        block_number = self.currentBlock().blockNumber()
        if block_number >= self.MAX_HIGHLIGHT_LINES:
            return

        # 先清除格式再应用
        self.setFormat(0, len(text), QTextCharFormat())
        for pattern, fmt in self.rules:
            for match in pattern.finditer(text):
                start, end = match.span()
                self.setFormat(start, end - start, fmt)


# ==================== 自定义编辑器 ====================
class CodeEditor(QTextEdit):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.file_path = None
        self.highlighter = None
        self._is_loading = False  # 加载标志，避免触发预览更新
        self.setFont(QFont("Consolas", 10))
        self.setTabStopDistance(4 * self.fontMetrics().horizontalAdvance(' '))

    def set_markdown_highlighting(self, enabled):
        if enabled:
            if self.highlighter is None:
                self.highlighter = MarkdownHighlighter(self.document())
            return
        if self.highlighter is not None:
            self.highlighter.setDocument(None)
            self.highlighter = None

    @staticmethod
    def is_markdown_path(path):
        ext = QFileInfo(path).suffix().lower()
        return ext in ["md", "markdown"]

    def set_file_path(self, path):
        self.file_path = path

    def load_file(self, path):
        try:
            with open(path, 'r', encoding='utf-8') as f:
                content = f.read()
            # 禁用高亮，使用批量编辑避免频繁触发
            self.set_markdown_highlighting(False)
            self._is_loading = True
            # 使用批量编辑提升大文件加载速度
            cursor = self.textCursor()
            cursor.beginEditBlock()
            cursor.select(QTextCursor.SelectionType.Document)
            cursor.removeSelectedText()
            cursor.insertText(content)
            cursor.endEditBlock()
            self.file_path = path
            self._is_loading = False
            # 延迟启用高亮，给UI一个喘息的机会
            QTimer.singleShot(50, lambda: self.set_markdown_highlighting(self.is_markdown_path(path)))
            self.document().setModified(False)
            return True
        except Exception as e:
            self._is_loading = False
            QMessageBox.critical(self, "错误", f"无法打开文件：{str(e)}")
            return False

    def save_to_path(self, path):
        try:
            with open(path, 'w', encoding='utf-8') as f:
                f.write(self.toPlainText())
            self.file_path = path
            self.set_markdown_highlighting(self.is_markdown_path(path))
            self.document().setModified(False)
            return True
        except Exception as e:
            QMessageBox.critical(self, "错误", f"无法保存文件：{str(e)}")
            return False


# ==================== 悬停放大容器（产品亮点，保留） ====================
class HoverWidget(QWidget):
    def __init__(self, widget, on_enter, on_leave, parent=None):
        super().__init__(parent)
        self.on_enter_callback = on_enter
        self.on_leave_callback = on_leave
        layout = QVBoxLayout()
        layout.setContentsMargins(0, 0, 0, 0)
        layout.addWidget(widget)
        self.setLayout(layout)
        self.widget = widget
        self.setAttribute(Qt.WidgetAttribute.WA_StyledBackground, True)

    def enterEvent(self, event):
        self.on_enter_callback(self)
        super().enterEvent(event)

    def leaveEvent(self, event):
        self.on_leave_callback(self)
        super().leaveEvent(event)


class PreviewBridge(QObject):
    scrollPercentChanged = pyqtSignal(float)

    @pyqtSlot(float)
    def reportScrollPercent(self, percent):
        self.scrollPercentChanged.emit(percent)


def get_resource_base_dir():
    return getattr(sys, "_MEIPASS", os.path.dirname(__file__))


def get_resource_url(*parts):
    return QUrl.fromLocalFile(os.path.join(get_resource_base_dir(), *parts))


def convert_special_markdown(text):
    """
    简化版 Markdown 扩展处理
    移除了 MathJax 和 Mermaid 的特殊处理
    仅保留下标/上标转换
    """
    # 下标 ~text~
    text = re.sub(r"(?<!\\)~([^~\n]+?)~", r"<sub>\1</sub>", text)
    # 上标 ^text^
    text = re.sub(r"(?<!\\)\^([^\^\n]+?)\^", r"<sup>\1</sup>", text)
    return text


def build_preview_html(markdown_text):
    """
    优化版 HTML 生成
    移除 MathJax 和 Mermaid，仅保留基础 Markdown 渲染
    """
    markdown_module = get_markdown_module()
    converted_text = convert_special_markdown(markdown_text)

    # 简化扩展：移除 codehilite（太重），保留基础功能
    body_html = markdown_module.markdown(
        converted_text,
        extensions=['extra', 'tables', 'fenced_code']
    )

    # 精简的 CSS 样式
    css = """
    <style>
        :root { color-scheme: light; }
        html, body {
            margin: 0; padding: 0;
            background: #f7f7fb;
            color: #1f2937;
            font-family: 'Segoe UI', 'Noto Sans SC', sans-serif;
            font-size: 14px; line-height: 1.75;
        }
        #content {
            max-width: 980px; margin: 0 auto;
            padding: 24px 28px 72px; box-sizing: border-box;
        }
        h1, h2, h3, h4, h5, h6 {
            line-height: 1.25; margin: 1.2em 0 0.6em; color: #0f172a;
        }
        h1 { font-size: 2rem; border-bottom: 1px solid #dbe1ea; padding-bottom: 0.3em; }
        h2 { font-size: 1.5rem; border-bottom: 1px solid #e5e7eb; padding-bottom: 0.2em; }
        h3 { font-size: 1.25rem; }
        p, ul, ol, blockquote, table, pre { margin: 0.9em 0; }
        a { color: #0366d6; text-decoration: none; }
        a:hover { text-decoration: underline; }
        code {
            font-family: 'Cascadia Mono', 'Consolas', monospace;
            background: #eef2f7; color: #0f172a;
            border-radius: 6px; padding: 0.15em 0.35em;
        }
        pre {
            background: #0b1020; color: #e5eefb;
            border-radius: 12px; padding: 16px; overflow: auto;
        }
        pre code { background: transparent; color: inherit; padding: 0; }
        blockquote {
            border-left: 4px solid #7c3aed; margin-left: 0;
            padding: 0.4em 1em; color: #4b5563;
            background: rgba(124, 58, 237, 0.06);
            border-radius: 0 10px 10px 0;
        }
        table { border-collapse: collapse; width: 100%; display: block; overflow-x: auto; }
        th, td { border: 1px solid #d1d5db; padding: 0.55em 0.8em; text-align: left; }
        th { background: #eef2ff; }
        img { max-width: 100%; height: auto; }
        hr { border: none; border-top: 1px solid #d1d5db; margin: 1.5em 0; }
    </style>
    """

    # 精简的 JavaScript：仅保留滚动同步，移除 MathJax/Mermaid
    script = """
    <script src="qrc:///qtwebchannel/qwebchannel.js"></script>
    <script>
        window.setScrollPercent = function(percent) {
            const doc = document.scrollingElement || document.documentElement;
            const maxScroll = Math.max(0, doc.scrollHeight - doc.clientHeight);
            const clamped = Math.max(0, Math.min(1, percent));
            doc.scrollTop = maxScroll <= 0 ? 0 : Math.round(maxScroll * clamped);
        };

        document.addEventListener('DOMContentLoaded', function() {
            new QWebChannel(qt.webChannelTransport, function(channel) {
                window.tmlBridge = channel.objects.tmlBridge;
                const sendScroll = function() {
                    const doc = document.scrollingElement || document.documentElement;
                    const maxScroll = Math.max(1, doc.scrollHeight - doc.clientHeight);
                    window.tmlBridge.reportScrollPercent(doc.scrollTop / maxScroll);
                };

                let scheduled = false;
                window.addEventListener('scroll', function() {
                    if (scheduled) return;
                    scheduled = true;
                    requestAnimationFrame(function() {
                        scheduled = false;
                        sendScroll();
                    });
                }, { passive: true });

                // 初始化滚动位置
                if (typeof window.__pendingScrollPercent === 'number') {
                    window.setScrollPercent(window.__pendingScrollPercent);
                }
            });
        });
    </script>
    """

    return f"""<!doctype html>
<html>
<head>
<meta charset="utf-8">
{css}
{script}
</head>
<body>
<div id="content">
{body_html}
</div>
</body>
</html>"""


# ==================== 主窗口 ====================
class MarkdownEditor(QMainWindow):
    def __init__(self, initial_paths=None):
        super().__init__()
        self.startup_paths = list(initial_paths or [])
        log_startup(f"window init begin, paths={len(self.startup_paths)}")
        icon_path = get_icon_path()
        if icon_path:
            self.setWindowIcon(QIcon(icon_path))
        self.setWindowTitle("TML Markdown 编辑器")
        self.resize(700, 600)
        self.setAcceptDrops(True)

        self.split_enabled = False
        self.hovered_side = None
        self._syncing_scroll = False
        self._preview_loaded = False
        self._preview_scroll_percent = 0.0

        # 防抖定时器：避免每次按键都触发预览更新
        self._preview_timer = QTimer()
        self._preview_timer.setSingleShot(True)
        self._preview_timer.timeout.connect(self._do_update_preview)

        # 中央分割器
        self.splitter = QSplitter(Qt.Orientation.Horizontal)
        self.setCentralWidget(self.splitter)

        # 左侧预览区（懒初始化）
        self.preview = None
        self.preview_bridge = None
        self.preview_channel = None
        preview_placeholder = QWidget()
        self.left_container = HoverWidget(
            preview_placeholder,
            self.on_hover_enter,
            self.on_hover_leave,
            self.splitter
        )

        # 右侧多标签页编辑区
        self.tab_widget = QTabWidget()
        self.tab_widget.setTabsClosable(True)
        self.tab_widget.tabCloseRequested.connect(self.close_tab)
        self.tab_widget.currentChanged.connect(self.on_current_tab_changed)
        self.right_container = HoverWidget(
            self.tab_widget,
            self.on_hover_enter,
            self.on_hover_leave,
            self.splitter
        )

        self.splitter.addWidget(self.left_container)
        self.splitter.addWidget(self.right_container)

        # 初始单屏模式
        self.left_container.hide()
        self.splitter.setSizes([0, self.width()])

        self.create_menu_bar()

        self.status_bar = QStatusBar()
        self.setStatusBar(self.status_bar)
        self.status_bar.showMessage("就绪")

        if self.startup_paths:
            QTimer.singleShot(0, self.open_startup_files)
        else:
            self.new_tab()
        log_startup("window init end")

    def open_startup_files(self):
        log_startup(f"opening startup files: {len(self.startup_paths)}")
        if not self.open_file_paths(self.startup_paths):
            self.new_tab()
        self.startup_paths = []

    def ensure_preview_initialized(self):
        if self.preview is not None:
            return

        log_startup("initializing preview webengine")
        from PyQt6.QtWebEngineWidgets import QWebEngineView
        from PyQt6.QtWebEngineCore import QWebEngineSettings
        from PyQt6.QtWebChannel import QWebChannel

        self.preview = QWebEngineView()
        self.preview.setContextMenuPolicy(Qt.ContextMenuPolicy.NoContextMenu)
        self.preview.settings().setAttribute(QWebEngineSettings.WebAttribute.LocalContentCanAccessFileUrls, True)
        self.preview.settings().setAttribute(QWebEngineSettings.WebAttribute.LocalContentCanAccessRemoteUrls, False)
        self.preview.page().loadFinished.connect(self.on_preview_load_finished)

        self.preview_bridge = PreviewBridge(self)
        self.preview_bridge.scrollPercentChanged.connect(self.on_preview_scroll_percent)
        self.preview_channel = QWebChannel(self.preview.page())
        self.preview_channel.registerObject("tmlBridge", self.preview_bridge)
        self.preview.page().setWebChannel(self.preview_channel)

        old_left = self.left_container
        self.left_container = HoverWidget(
            self.preview,
            self.on_hover_enter,
            self.on_hover_leave,
            self.splitter
        )
        self.splitter.insertWidget(0, self.left_container)
        old_left.hide()
        old_left.setParent(None)
        old_left.deleteLater()

    # ========== 滚动同步 ==========
    def sync_scroll(self, source_scrollbar, target_scrollbar):
        if self._syncing_scroll or not self.split_enabled:
            return
        self._syncing_scroll = True
        src_min = source_scrollbar.minimum()
        src_max = source_scrollbar.maximum()
        src_val = source_scrollbar.value()
        if src_max > src_min:
            percent = (src_val - src_min) / (src_max - src_min)
            tgt_min = target_scrollbar.minimum()
            tgt_max = target_scrollbar.maximum()
            tgt_val = int(tgt_min + percent * (tgt_max - tgt_min))
            target_scrollbar.setValue(tgt_val)
        self._syncing_scroll = False

    def scrollbar_percent(self, scrollbar):
        minimum = scrollbar.minimum()
        maximum = scrollbar.maximum()
        if maximum <= minimum:
            return 0.0
        return (scrollbar.value() - minimum) / (maximum - minimum)

    def set_scrollbar_percent(self, scrollbar, percent):
        minimum = scrollbar.minimum()
        maximum = scrollbar.maximum()
        if maximum <= minimum:
            scrollbar.setValue(minimum)
            return
        clamped = max(0.0, min(1.0, percent))
        scrollbar.setValue(int(minimum + clamped * (maximum - minimum)))

    def set_preview_scroll_percent(self, percent):
        self._preview_scroll_percent = max(0.0, min(1.0, percent))
        if not self._preview_loaded or not self.split_enabled or self.preview is None:
            return
        self.preview.page().runJavaScript(f"window.setScrollPercent({self._preview_scroll_percent:.8f});")

    def connect_scroll_sync(self):
        editor = self.current_editor()
        if editor and self.split_enabled:
            editor_vscroll = editor.verticalScrollBar()
            try:
                editor_vscroll.valueChanged.disconnect(self.on_editor_scroll)
            except TypeError:
                pass
            editor_vscroll.valueChanged.connect(self.on_editor_scroll)
            self.on_editor_scroll(editor_vscroll.value())

    def disconnect_scroll_sync(self):
        editor = self.current_editor()
        if editor:
            try:
                editor.verticalScrollBar().valueChanged.disconnect(self.on_editor_scroll)
            except TypeError:
                pass

    def on_editor_scroll(self, value):
        if not self.split_enabled or self._syncing_scroll:
            return
        editor = self.current_editor()
        if editor:
            percent = self.scrollbar_percent(editor.verticalScrollBar())
            self.set_preview_scroll_percent(percent)

    def on_preview_scroll_percent(self, percent):
        if not self.split_enabled or self._syncing_scroll:
            return
        editor = self.current_editor()
        if editor:
            current_percent = self.scrollbar_percent(editor.verticalScrollBar())
            if abs(current_percent - percent) < 0.01:
                return
            self._syncing_scroll = True
            self.set_scrollbar_percent(editor.verticalScrollBar(), percent)
            self._syncing_scroll = False

    # ========== 悬停放大（产品亮点，保留） ==========
    def on_hover_enter(self, hover_widget):
        if not self.split_enabled:
            return
        if hover_widget == self.left_container:
            self.hovered_side = "left"
        elif hover_widget == self.right_container:
            self.hovered_side = "right"
        self.adjust_splitter_sizes()

    def on_hover_leave(self, hover_widget):
        if not self.split_enabled:
            return
        self.hovered_side = None
        self.adjust_splitter_sizes()

    def adjust_splitter_sizes(self):
        if not self.split_enabled:
            return
        total = self.splitter.width()
        if total <= 0:
            return
        if self.hovered_side == "left":
            self.splitter.setSizes([int(total * 0.7), int(total * 0.3)])
        elif self.hovered_side == "right":
            self.splitter.setSizes([int(total * 0.3), int(total * 0.7)])
        else:
            self.splitter.setSizes([total // 2, total // 2])

    def resizeEvent(self, event):
        super().resizeEvent(event)
        if self.split_enabled:
            self.adjust_splitter_sizes()

    # ========== 分屏模式 ==========
    def set_split_mode(self, enabled):
        self.split_enabled = enabled
        if enabled:
            self.ensure_preview_initialized()
            self.left_container.show()
            total = self.splitter.width()
            if total > 0:
                self.splitter.setSizes([total // 2, total // 2])
            self.update_preview()
            self.connect_scroll_sync()
            self.status_bar.showMessage("分屏模式已开启")
        else:
            self.disconnect_scroll_sync()
            self.left_container.hide()
            self.splitter.setSizes([0, self.splitter.width()])
            self.status_bar.showMessage("分屏模式已关闭")
        self.split_action.setChecked(enabled)

    def toggle_split_mode(self):
        self.set_split_mode(not self.split_enabled)

    # ========== 多标签页管理 ==========
    def new_tab(self, file_path=None, content=""):
        editor = CodeEditor()
        if file_path:
            if not editor.load_file(file_path):
                return None
            tab_name = os.path.basename(file_path)
        else:
            editor.setPlainText(content)
            tab_name = "未命名"
        index = self.tab_widget.addTab(editor, tab_name)
        self.tab_widget.setCurrentIndex(index)
        editor.textChanged.connect(lambda: self.on_editor_text_changed(editor))
        editor.document().modificationChanged.connect(
            lambda modified: self.update_tab_title(editor, modified)
        )
        if self.split_enabled:
            self.connect_scroll_sync()
        return editor

    def open_file_paths(self, paths):
        valid_paths = [path for path in paths if path and os.path.isfile(path)]
        if not valid_paths:
            return False

        log_startup(f"open_file_paths valid={len(valid_paths)}")
        opened_any = False
        for path in valid_paths:
            log_startup(f"opening file: {os.path.basename(path)}")
            editor = self.new_tab(file_path=path)
            if editor is None:
                log_startup(f"failed file: {path}")
                continue
            opened_any = True
            self.status_bar.showMessage(f"已打开：{path}")

        if opened_any:
            current = self.current_editor()
            if current and current.file_path:
                ext = QFileInfo(current.file_path).suffix().lower()
                is_md = ext in ["md", "markdown"]
                if is_md != self.split_enabled:
                    self.set_split_mode(is_md)
                elif self.split_enabled:
                    self.update_preview()

        return opened_any

    def close_tab(self, index):
        editor = self.tab_widget.widget(index)
        if editor and editor.document().isModified():
            tab_name = self.tab_widget.tabText(index)
            ret = QMessageBox.question(
                self, "未保存",
                f"文档「{tab_name}」已修改，是否保存？",
                QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No | QMessageBox.StandardButton.Cancel
            )
            if ret == QMessageBox.StandardButton.Yes:
                if not self.save_current_editor(editor):
                    return
            elif ret == QMessageBox.StandardButton.Cancel:
                return
        self.tab_widget.removeTab(index)
        if self.tab_widget.count() == 0:
            self.new_tab()
        elif self.split_enabled:
            self.connect_scroll_sync()

    def update_tab_title(self, editor, modified):
        index = self.tab_widget.indexOf(editor)
        if index == -1:
            return
        base_name = os.path.basename(editor.file_path) if editor.file_path else "未命名"
        title = base_name + " *" if modified else base_name
        self.tab_widget.setTabText(index, title)

    def on_editor_text_changed(self, editor):
        """使用防抖机制延迟预览更新，避免每次按键都触发"""
        # 如果正在加载文件，不触发预览更新
        if getattr(editor, '_is_loading', False):
            return
        if self.split_enabled and editor == self.current_editor():
            self._preview_timer.start(300)  # 300ms 防抖

    def _do_update_preview(self):
        """实际执行预览更新"""
        editor = self.current_editor()
        if not editor or not self.split_enabled:
            return
        self.ensure_preview_initialized()
        log_startup("rendering preview")
        md_text = editor.toPlainText()
        editor_vscroll = editor.verticalScrollBar()
        self._preview_scroll_percent = self.scrollbar_percent(editor_vscroll)
        self._preview_loaded = False
        try:
            full_html = build_preview_html(md_text)
        except Exception as e:
            err_msg = html_lib.escape(str(e))
            full_html = f"""<!doctype html>
<html>
<head>
<meta charset="utf-8">
<style>
body {{ background: #f7f7fb; color: #1f2937; font-family: sans-serif; padding: 24px; }}
h1 {{ color: #dc2626; }}
pre {{ background: #fee2e2; padding: 12px; border-radius: 6px; overflow-x: auto; }}
</style>
</head>
<body>
<h1>渲染错误</h1>
<p>Markdown 渲染失败，请检查依赖是否正确安装。</p>
<pre>{err_msg}</pre>
</body>
</html>"""
        self.preview.setHtml(full_html)

    def on_current_tab_changed(self, index):
        editor = self.current_editor()
        if editor:
            path = editor.file_path
            if path:
                self.setWindowTitle(f"TML Markdown 编辑器 - {os.path.basename(path)}")
                ext = QFileInfo(path).suffix().lower()
                is_md = ext in ["md", "markdown"]
                if is_md != self.split_enabled:
                    self.set_split_mode(is_md)
            else:
                self.setWindowTitle("TML Markdown 编辑器")
            if self.split_enabled:
                self.connect_scroll_sync()
            self.status_bar.showMessage(f"当前文件：{path if path else '未保存'}")

    def current_editor(self):
        return self.tab_widget.currentWidget()

    def save_current_editor(self, editor=None):
        if editor is None:
            editor = self.current_editor()
        if not editor:
            return False
        if editor.file_path:
            return editor.save_to_path(editor.file_path)
        else:
            return self.save_as_current_editor(editor)

    def save_as_current_editor(self, editor=None):
        if editor is None:
            editor = self.current_editor()
        if not editor:
            return False
        path, _ = QFileDialog.getSaveFileName(
            self, "保存文件", "",
            "Markdown 文件 (*.md);;文本文件 (*.txt);;LaTeX 文件 (*.tex);;所有文件 (*)"
        )
        if path:
            if editor.save_to_path(path):
                self.update_tab_title(editor, False)
                ext = QFileInfo(path).suffix().lower()
                if ext in ["md", "markdown"] and not self.split_enabled:
                    reply = QMessageBox.question(self, "建议", "是否开启分屏模式以预览 Markdown？",
                                                 QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No)
                    if reply == QMessageBox.StandardButton.Yes:
                        self.set_split_mode(True)
                return True
        return False

    # ========== 文件操作 ==========
    def new_file(self):
        self.new_tab()

    def open_file(self):
        path, _ = QFileDialog.getOpenFileName(
            self, "打开文件", "",
            "文本文件 (*.txt *.md *.markdown *.tex);;所有文件 (*)"
        )
        if path:
            self.open_file_paths([path])

    def dragEnterEvent(self, event):
        if event.mimeData().hasUrls() and any(url.isLocalFile() for url in event.mimeData().urls()):
            event.acceptProposedAction()
        else:
            super().dragEnterEvent(event)

    def dropEvent(self, event):
        paths = [url.toLocalFile() for url in event.mimeData().urls() if url.isLocalFile()]
        if self.open_file_paths(paths):
            event.acceptProposedAction()
        else:
            super().dropEvent(event)

    def save_file(self):
        editor = self.current_editor()
        if editor:
            if editor.file_path:
                editor.save_to_path(editor.file_path)
                self.update_tab_title(editor, False)
            else:
                self.save_as_current_editor(editor)

    def save_as_file(self):
        self.save_as_current_editor()

    # ========== 预览更新（直接调用，无需防抖） ==========
    def update_preview(self):
        """立即更新预览（用于切换标签等场景）"""
        self._preview_timer.stop()
        self._do_update_preview()

    def on_preview_load_finished(self, ok):
        self._preview_loaded = bool(ok)
        if not ok or not self.split_enabled:
            return
        self.preview.page().runJavaScript(
            f"window.__pendingScrollPercent = {self._preview_scroll_percent:.8f};"
        )

    # ========== 导出 Word ==========
    def export_to_word(self):
        editor = self.current_editor()
        if not editor:
            return
        Document = get_docx_document_class()
        if Document is None:
            QMessageBox.warning(self, "缺少依赖", "未安装 python-docx，无法导出 Word。")
            return
        path, _ = QFileDialog.getSaveFileName(
            self, "导出为 Word", "", "Word 文档 (*.docx)"
        )
        if not path:
            return
        try:
            doc = Document()
            self.add_markdown_to_docx(doc, editor.toPlainText())
            doc.save(path)
            self.status_bar.showMessage(f"已导出：{path}")
        except Exception as e:
            QMessageBox.critical(self, "错误", f"导出失败：{str(e)}")

    def add_markdown_to_docx(self, doc, text):
        import re
        lines = text.splitlines()
        in_code_block = False
        for line in lines:
            stripped = line.rstrip()
            if stripped.startswith("```"):
                in_code_block = not in_code_block
                continue
            if in_code_block:
                paragraph = doc.add_paragraph()
                run = paragraph.add_run(stripped)
                run.font.name = "Consolas"
                continue
            if not stripped:
                doc.add_paragraph("")
                continue
            if re.match(r"^#{1,6}\s+", stripped):
                level = len(stripped.split(" ", 1)[0])
                title = stripped[level + 1:]
                doc.add_heading(title, level=level)
                continue
            if re.match(r"^\d+\.\s+", stripped):
                content = re.sub(r"^\d+\.\s+", "", stripped)
                paragraph = doc.add_paragraph(style="List Number")
                self.add_inline_runs(paragraph, content)
                continue
            if stripped.startswith("- ") or stripped.startswith("* "):
                content = stripped[2:]
                paragraph = doc.add_paragraph(style="List Bullet")
                self.add_inline_runs(paragraph, content)
                continue
            if stripped.startswith("> "):
                paragraph = doc.add_paragraph(style="Intense Quote")
                self.add_inline_runs(paragraph, stripped[2:])
                continue
            paragraph = doc.add_paragraph()
            self.add_inline_runs(paragraph, stripped)

    def add_inline_runs(self, paragraph, text):
        import re
        token_re = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)")
        pos = 0
        for match in token_re.finditer(text):
            if match.start() > pos:
                paragraph.add_run(text[pos:match.start()])
            token = match.group(0)
            if token.startswith("**"):
                run = paragraph.add_run(token[2:-2])
                run.bold = True
            elif token.startswith("*"):
                run = paragraph.add_run(token[1:-1])
                run.italic = True
            elif token.startswith("`"):
                run = paragraph.add_run(token[1:-1])
                run.font.name = "Consolas"
            pos = match.end()
        if pos < len(text):
            paragraph.add_run(text[pos:])

    # ========== 帮助窗口 ==========
    def show_help(self):
        dialog = QDialog(self)
        dialog.setWindowTitle("使用指导")
        dialog.resize(520, 420)

        layout = QVBoxLayout()
        guide = QTextEdit()
        guide.setReadOnly(True)
        guide.setPlainText(
            "使用指导\n"
            "\n"
            "1. 新建/打开\n"
            "- 文件菜单中可新建或打开文件。\n"
            "- 支持.txt, .md, .tex, .py文件。\n"
            "\n"
            "2. 分屏预览\n"
            "- 打开.md类型文件自动进入分屏模式。\n"
            "- 视图菜单可开启或关闭分屏模式。\n"
            "- 分屏时左侧预览，右侧编辑，滚动自动同步。\n"
            "- 鼠标悬停在某侧可自动放大该侧比例。\n"
            "\n"
            "3. 预览缩放\n"
            "- 视图菜单可放大或缩小预览字体。\n"
            "\n"
            "4. 导出 Word\n"
            "- 文件菜单可导出为 .docx。\n"
            "\n"
            "本项目旨在提供一个方便、轻量的小窗口Markdown编辑体验\n"
            "欢迎反馈和建议！\n"
            "问题提交：https://github.com/Loryage\n"
        )
        layout.addWidget(guide)

        close_btn = QPushButton("关闭")
        close_btn.clicked.connect(dialog.accept)
        layout.addWidget(close_btn)

        dialog.setLayout(layout)
        dialog.exec()

    # ========== 界面构建 ==========
    def create_menu_bar(self):
        menubar = self.menuBar()
        file_menu = menubar.addMenu("文件(&F)")
        new_action = QAction("新建(&N)", self)
        new_action.triggered.connect(self.new_file)
        new_action.setShortcut("Ctrl+N")
        file_menu.addAction(new_action)

        open_action = QAction("打开(&O)...", self)
        open_action.triggered.connect(self.open_file)
        open_action.setShortcut("Ctrl+O")
        file_menu.addAction(open_action)

        save_action = QAction("保存(&S)", self)
        save_action.triggered.connect(self.save_file)
        save_action.setShortcut("Ctrl+S")
        file_menu.addAction(save_action)

        save_as_action = QAction("另存为(&A)...", self)
        save_as_action.triggered.connect(self.save_as_file)
        file_menu.addAction(save_as_action)

        export_action = QAction("导出为 Word(&W)...", self)
        export_action.triggered.connect(self.export_to_word)
        file_menu.addAction(export_action)

        file_menu.addSeparator()
        exit_action = QAction("退出(&X)", self)
        exit_action.triggered.connect(self.close)
        file_menu.addAction(exit_action)

        view_menu = menubar.addMenu("视图(&V)")
        zoom_in_action = QAction("放大预览(&+)", self)
        zoom_in_action.triggered.connect(lambda: self.preview_zoom_in())
        view_menu.addAction(zoom_in_action)

        zoom_out_action = QAction("缩小预览(&-)", self)
        zoom_out_action.triggered.connect(lambda: self.preview_zoom_out())
        view_menu.addAction(zoom_out_action)

        view_menu.addSeparator()
        self.split_action = QAction("分屏模式(&P)", self)
        self.split_action.setCheckable(True)
        self.split_action.setChecked(False)
        self.split_action.triggered.connect(self.toggle_split_mode)
        view_menu.addAction(self.split_action)

        help_menu = menubar.addMenu("帮助(&H)")
        help_action = QAction("使用指导(&G)", self)
        help_action.triggered.connect(self.show_help)
        help_menu.addAction(help_action)

    def preview_zoom_in(self):
        if self.split_enabled and self.preview is not None:
            self.preview.setZoomFactor(self.preview.zoomFactor() + 0.1)

    def preview_zoom_out(self):
        if self.split_enabled and self.preview is not None:
            self.preview.setZoomFactor(max(0.4, self.preview.zoomFactor() - 0.1))


def get_icon_path():
    base_dir = getattr(sys, "_MEIPASS", os.path.dirname(__file__))
    icon_path = os.path.join(base_dir, "tml.ico")
    return icon_path if os.path.exists(icon_path) else None


def set_windows_app_id(app_id):
    try:
        ctypes.windll.shell32.SetCurrentProcessExplicitAppUserModelID(app_id)
    except Exception:
        pass


if __name__ == "__main__":
    log_startup("app entry")
    install_global_exception_hook()
    set_windows_app_id("TML.TMLEditor")
    log_startup("set app id done")

    try:
        import pyi_splash
        _has_pyi_splash = True
    except ImportError:
        _has_pyi_splash = False

    QApplication.setAttribute(Qt.ApplicationAttribute.AA_ShareOpenGLContexts, True)
    log_startup("AA_ShareOpenGLContexts enabled")
    app = QApplication(sys.argv)
    log_startup("QApplication created")
    icon_path = get_icon_path()
    if icon_path:
        app.setWindowIcon(QIcon(icon_path))
    log_startup("window icon applied" if icon_path else "window icon skipped")
    initial_paths = [path for path in sys.argv[1:] if os.path.isfile(path)]
    log_startup(f"argv files={len(initial_paths)}")

    splash = SplashScreen()
    splash.show()
    app.processEvents()
    log_startup("splash shown")
    splash.set_progress(10, "正在初始化界面")
    app.processEvents()

    window = None

    def close_pyi_splash():
        global _has_pyi_splash
        if _has_pyi_splash:
            try:
                import pyi_splash
                if pyi_splash.is_alive():
                    pyi_splash.close()
                    log_startup("pyi_splash closed")
            except Exception:
                pass
            _has_pyi_splash = False

    def finish_startup():
        global window
        splash.set_progress(100, "启动完成")
        app.processEvents()
        close_pyi_splash()
        QTimer.singleShot(200, _final_close)

    def _final_close():
        global window
        try:
            splash.finish(window)
        except Exception:
            splash.close()
        log_startup("splash finished")

    def step_create_window():
        global window
        splash.set_progress(40, "创建主窗口")
        app.processEvents()
        window = MarkdownEditor(initial_paths)
        log_startup("window created")
        splash.set_progress(75, "加载组件")
        app.processEvents()
        QTimer.singleShot(10, step_show_window)

    def step_show_window():
        splash.set_progress(95, "准备就绪")
        app.processEvents()
        window.show()
        window.raise_()
        window.activateWindow()
        log_startup("window shown")
        QTimer.singleShot(50, finish_startup)

    QTimer.singleShot(10, step_create_window)
    sys.exit(app.exec())
